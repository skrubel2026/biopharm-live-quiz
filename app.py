"""
Biopharmaceutics Live Quiz — Streamlit edition
================================================
A Kahoot-style timed quiz for classroom use.

HOW STUDENTS/HOST CONNECT
--------------------------
This one app serves two roles based on a URL parameter:
  - Host (classroom screen):  https://<your-app>.streamlit.app/?role=host
  - Student (on phones):      https://<your-app>.streamlit.app/?role=student

Share the STUDENT link with your class (write it on the board, paste it in
your class group chat, etc). Open the HOST link yourself on the classroom
computer/projector.

EDITING QUESTIONS
------------------
Edit the QUESTIONS list below. Each question needs:
  q             - the question text
  options       - list of exactly 4 answer choices
  correct_index - which option (0-3) is correct
  time          - seconds allowed for that question

No other code needs to change when you edit questions.
"""

import time
import re
import uuid
import io
import csv
import datetime
import streamlit as st
from streamlit_autorefresh import st_autorefresh
from docx import Document

try:
    import gspread
    from google.oauth2.service_account import Credentials
    GSPREAD_AVAILABLE = True
except Exception:
    GSPREAD_AVAILABLE = False

# ============================================================
# QUESTION BANK — edit freely
# ============================================================
QUESTIONS = [
    {"q": "Which process best describes the release of drug from its dosage form, making it available for absorption?",
     "options": ["Disintegration", "Dissolution", "Liberation", "Elimination"], "correct_index": 2, "time": 30},
    {"q": "Bioavailability (F) of a drug given by IV bolus is assumed to be:",
     "options": ["0%", "50%", "75%", "100%"], "correct_index": 3, "time": 30},
    {"q": "Which factor does NOT typically affect the rate of drug dissolution?",
     "options": ["Particle size", "Drug pKa/solubility", "Patient's blood type", "Agitation/hydrodynamics"], "correct_index": 2, "time": 30},
    {"q": "First-pass metabolism primarily reduces the bioavailability of drugs administered:",
     "options": ["Intravenously", "Orally", "Sublingually", "Transdermally (in most cases)"], "correct_index": 1, "time": 25},
    {"q": "According to the Noyes-Whitney equation, dissolution rate is directly proportional to:",
     "options": ["Drug's melting point", "Surface area of the particle", "Patient's body weight", "Gastric emptying time"], "correct_index": 1, "time": 30},
    {"q": "A drug with high permeability and low solubility falls into which BCS class?",
     "options": ["Class I", "Class II", "Class III", "Class IV"], "correct_index": 1, "time": 25},
    {"q": "Which parameter describes the fraction of an administered dose that reaches systemic circulation unchanged?",
     "options": ["Clearance", "Volume of distribution", "Bioavailability", "Half-life"], "correct_index": 2, "time": 25},
    {"q": "Enteric coating on a tablet is primarily designed to:",
     "options": ["Speed up gastric dissolution", "Protect the drug from stomach acid or protect stomach from drug", "Improve taste only", "Increase tablet hardness"], "correct_index": 1, "time": 25},
    {"q": "In bioequivalence studies, two products are generally considered bioequivalent if the 90% CI of the AUC and Cmax ratio falls within:",
     "options": ["50-150%", "70-130%", "80-125%", "90-110%"], "correct_index": 2, "time": 30},
    {"q": "Which route of administration avoids first-pass hepatic metabolism entirely?",
     "options": ["Oral", "Rectal (upper)", "Intravenous", "Buccal (partially, but IV is complete)"], "correct_index": 2, "time": 25},
]
REVEAL_SECONDS = 6  # how long the answer reveal / round leaderboard shows before auto-advancing

# ============================================================
# GOOGLE SHEETS — question storage that survives app restarts
# ============================================================
# 1. Paste your Google Sheet's URL below (the sheet needs one tab named
#    exactly "Questions" with this header row in row 1:
#    Question | OptionA | OptionB | OptionC | OptionD | Correct | Time
# 2. Share that sheet (Editor access) with your service account's email.
# 3. Put the service account's JSON key into Streamlit Cloud -> your app ->
#    Settings -> Secrets, under the key name  gcp_service_account
#    (see the setup guide for the exact format).
SHEET_URL = "https://docs.google.com/spreadsheets/d/1sXwJWDIUSc753JwnMURkOSEabU45WlbdV3_71UgipHY/edit?gid=0#gid=0"  # <-- change this
WORKSHEET_NAME = "Questions"


def _get_gsheet_client():
    if not GSPREAD_AVAILABLE:
        return None, "The gspread library isn't installed (check requirements.txt)."
    try:
        creds_dict = dict(st.secrets["gcp_service_account"])
    except Exception:
        return None, "No Google service account found in Streamlit Secrets (gcp_service_account)."
    try:
        scopes = ["https://www.googleapis.com/auth/spreadsheets", "https://www.googleapis.com/auth/drive"]
        creds = Credentials.from_service_account_info(creds_dict, scopes=scopes)
        return gspread.authorize(creds), None
    except Exception as e:
        return None, f"Couldn't authenticate with Google: {e}"


def load_questions_from_sheet():
    """Returns (questions_or_None, error_message_or_None)."""
    if not SHEET_URL or "PASTE_YOUR" in SHEET_URL:
        return None, "No Google Sheet URL set yet (SHEET_URL in app.py)."
    client, err = _get_gsheet_client()
    if client is None:
        return None, err
    try:
        sh = client.open_by_url(SHEET_URL)
        ws = sh.worksheet(WORKSHEET_NAME)
        records = ws.get_all_records()
    except Exception as e:
        return None, f"Couldn't read the Google Sheet: {e}"

    questions = []
    for r in records:
        try:
            q_text = str(r.get("Question", "")).strip()
            correct_letter = str(r.get("Correct", "")).strip().upper()
            opts = [str(r.get("OptionA", "")).strip(), str(r.get("OptionB", "")).strip(),
                    str(r.get("OptionC", "")).strip(), str(r.get("OptionD", "")).strip()]
            if not q_text or correct_letter not in ("A", "B", "C", "D") or any(not o for o in opts):
                continue  # skip blank/incomplete rows rather than treating them as a question
            questions.append({
                "q": q_text,
                "options": opts,
                "correct_index": "ABCD".index(correct_letter),
                "time": int(r["Time"]) if str(r.get("Time", "")).strip() else 30,
            })
        except Exception:
            continue  # skip any malformed row rather than failing the whole load
    if not questions:
        return None, "Connected to the sheet, but no valid question rows were found."
    return questions, None


def save_questions_to_sheet(questions):
    """Returns (success_bool, error_message_or_None)."""
    if not SHEET_URL or "PASTE_YOUR" in SHEET_URL:
        return False, "No Google Sheet URL set yet (SHEET_URL in app.py)."
    client, err = _get_gsheet_client()
    if client is None:
        return False, err
    try:
        sh = client.open_by_url(SHEET_URL)
        ws = sh.worksheet(WORKSHEET_NAME)
        header = ["Question", "OptionA", "OptionB", "OptionC", "OptionD", "Correct", "Time"]
        rows = [header]
        for q in questions:
            rows.append([q["q"], q["options"][0], q["options"][1], q["options"][2], q["options"][3], "ABCD"[q["correct_index"]], q["time"]])
        ws.clear()
        ws.update(values=rows, range_name="A1")
        return True, None
    except Exception as e:
        return False, f"Couldn't save to the Google Sheet: {e}"


# ============================================================
# SHARED STATE (one instance shared by every visitor to this app)
# ============================================================
@st.cache_resource
def get_state():
    sheet_questions, sheet_error = load_questions_from_sheet()
    return {
        "status": "lobby",          # lobby | active | reveal | finished
        "current_q": -1,
        "question_started_at": 0.0,
        "reveal_started_at": 0.0,
        "questions": sheet_questions if sheet_questions else [dict(q) for q in QUESTIONS],
        "using_sheet": sheet_questions is not None,
        "sheet_error": sheet_error,
        "roster": {},                # pid -> {"name": str, "sid": str}
        "answers": {},                # q_index -> {pid: {"choice", "correct", "elapsed", "score"}}
    }


def rank_and_score_round(state, q_idx, time_limit):
    """Called once, right when a question's timer ends. Ranks every correct
    answer for that question by how fast it was (ties within 0.1s count as
    the same rank and get the same score), then writes a 'score' into each
    answer record. Fastest correct = 100, each following distinct rank loses
    10, floor of 10 points for any correct answer. Wrong/no answer = 0."""
    qanswers = state["answers"].get(q_idx, {})
    correct_entries = sorted(
        [(pid, a["elapsed"]) for pid, a in qanswers.items() if a.get("correct")],
        key=lambda x: x[1],
    )
    rank = 0
    last_time = None
    for pid, elapsed in correct_entries:
        if last_time is None or round(elapsed, 1) > round(last_time, 1):
            rank += 1
            last_time = elapsed
        score = max(10, 100 - (rank - 1) * 10)
        qanswers[pid]["score"] = score
    for pid, a in qanswers.items():
        if not a.get("correct"):
            a["score"] = 0


def leaderboard(state):
    totals = {pid: 0 for pid in state["roster"]}
    for qdict in state["answers"].values():
        for pid, a in qdict.items():
            totals[pid] = totals.get(pid, 0) + a.get("score", 0)
    rows = []
    for pid, s in totals.items():
        info = state["roster"].get(pid, {"name": "?", "sid": "?"})
        rows.append({"pid": pid, "name": info["name"], "sid": info["sid"], "score": s})
    rows.sort(key=lambda r: r["score"], reverse=True)
    return rows


def my_progress(state, pid, up_to_q):
    """How many questions this student has gotten correct so far, out of
    how many they've attempted (only counts questions up to and including
    the current one, so it never reveals future questions)."""
    correct_count, attempted = 0, 0
    for qi in range(0, up_to_q + 1):
        a = state["answers"].get(qi, {}).get(pid)
        if a:
            attempted += 1
            if a.get("correct"):
                correct_count += 1
    return correct_count, attempted


def detailed_results(state):
    """One row per student per question — exactly what they selected, whether
    it was correct, how long they took, and points earned. This is the audit
    trail: it comes straight from what was recorded live during the quiz,
    not from anything re-entered afterward."""
    rows = []
    for pid, info in state["roster"].items():
        for qi, q in enumerate(state["questions"]):
            a = state["answers"].get(qi, {}).get(pid)
            correct_letter = chr(65 + q["correct_index"])
            correct_text = q["options"][q["correct_index"]]
            if a is None:
                selected = "No answer recorded"
                was_correct = "No"
                took = ""
                pts = 0
            elif a["choice"] == -1:
                selected = "No answer (timed out)"
                was_correct = "No"
                took = round(a.get("elapsed", 0), 2)
                pts = a.get("score", 0)
            else:
                selected = f"{chr(65+a['choice'])}. {q['options'][a['choice']]}"
                was_correct = "Yes" if a.get("correct") else "No"
                took = round(a.get("elapsed", 0), 2)
                pts = a.get("score", 0)
            rows.append({
                "Student Name": info["name"],
                "Student ID": info["sid"],
                "Question #": qi + 1,
                "Question": q["q"],
                "Selected Answer": selected,
                "Correct Answer": f"{correct_letter}. {correct_text}",
                "Was Correct": was_correct,
                "Time Taken (s)": took,
                "Points Earned": pts,
            })
    return rows


def rows_to_csv_bytes(rows):
    if not rows:
        return b""
    buf = io.StringIO()
    writer = csv.DictWriter(buf, fieldnames=list(rows[0].keys()))
    writer.writeheader()
    writer.writerows(rows)
    return buf.getvalue().encode("utf-8")


HOST_PASSWORD = "biopharm2026"  # <-- change this any time: edit this line on GitHub, commit, and it takes effect on redeploy


def require_host_password(state):
    """Gate everything behind this. Returns True once the correct password
    has been entered in THIS browser tab/session — a different device or
    tab (e.g. a student who stumbles onto the host link) always has to
    enter it fresh, regardless of what's happening in the live quiz."""
    if st.session_state.get("host_authed"):
        return True
    st.markdown('<div class="eyebrow">HOST · BIOPHARMACEUTICS LIVE QUIZ</div>', unsafe_allow_html=True)
    st.title("🔒 Host access")
    st.write("This screen is for instructors only. Enter the host password to continue.")
    pw = st.text_input("Password", type="password", key="host_pw_input")
    if st.button("Unlock", type="primary"):
        if pw == HOST_PASSWORD:
            st.session_state.host_authed = True
            st.rerun()
        else:
            st.error("Incorrect password.")
    return False


def reset_quiz(state):
    state["status"] = "lobby"
    state["current_q"] = -1
    state["question_started_at"] = 0.0
    state["reveal_started_at"] = 0.0
    state["roster"] = {}
    state["answers"] = {}


# ============================================================
# QUESTION FILE PARSING (Word .docx or plain .txt upload)
# ============================================================
# Expected format, one block per question, blank line between blocks:
#
#   Q: Which process best describes drug release from its dosage form?
#   A) Disintegration
#   B) Dissolution
#   C) Liberation
#   D) Elimination
#   Correct: C
#   Time: 30
#
# "Time:" is optional (defaults to 30s). "Q1:", "Q2:" etc. also work.
def parse_questions_text(paragraphs):
    q_pattern = re.compile(r'^Q\s*\d*\s*[:.\)]\s*(.+)$', re.IGNORECASE)
    opt_pattern = re.compile(r'^([A-D])\s*[).:\-]\s*(.+)$', re.IGNORECASE)
    correct_pattern = re.compile(r'^Correct\s*[:.\-]?\s*([A-D])', re.IGNORECASE)
    time_pattern = re.compile(r'^Time\s*[:.\-]?\s*(\d+)', re.IGNORECASE)

    questions = []
    current = None
    errors = []

    def flush(label):
        nonlocal current
        if current is None:
            return
        opts = current.get("options", {})
        missing = [L for L in "ABCD" if L not in opts]
        preview = (current.get("q") or "")[:50]
        if not current.get("q"):
            errors.append(f"A question block near '{label}' has no question text.")
        elif missing:
            errors.append(f"'{preview}...' is missing option(s): {', '.join(missing)}.")
        elif current.get("correct") is None:
            errors.append(f"'{preview}...' has no 'Correct:' line.")
        else:
            questions.append({
                "q": current["q"],
                "options": [opts["A"], opts["B"], opts["C"], opts["D"]],
                "correct_index": "ABCD".index(current["correct"]),
                "time": current.get("time", 30),
            })
        current = None

    for raw in paragraphs:
        line = raw.strip()
        if not line:
            continue
        m = q_pattern.match(line)
        if m:
            flush(line)
            current = {"q": m.group(1).strip(), "options": {}}
            continue
        if current is None:
            continue  # ignore stray text before the first "Q:" line
        m = opt_pattern.match(line)
        if m:
            current["options"][m.group(1).upper()] = m.group(2).strip()
            continue
        m = correct_pattern.match(line)
        if m:
            current["correct"] = m.group(1).upper()
            continue
        m = time_pattern.match(line)
        if m:
            current["time"] = int(m.group(1))
            continue
    flush("end of file")
    return questions, errors


def parse_docx_table(table):
    """Parses the 'numbered table' quiz format: one column for the question
    number, two columns holding the stem / option pairs (A+B on one row,
    C+D on the next), and a final column repeating the correct answer letter
    on every row belonging to that question."""
    groups = {}
    order = []
    for r in table.rows:
        cells = [c.text.strip() for c in r.cells]
        if len(cells) < 3:
            continue
        num = cells[0].strip()
        if not num or not num.replace(".", "").isdigit():
            continue  # header row or a stray row with no question number
        if num not in groups:
            groups[num] = []
            order.append(num)
        groups[num].append(cells)

    opt_re = re.compile(r'^([A-D])\s*[.\)]\s*(.*)$', re.IGNORECASE)
    questions, errors = [], []
    for num in order:
        stem = None
        opts = {}
        correct = None
        for cells in groups[num]:
            last = cells[-1].strip()
            if last and len(last) <= 2 and last.upper() in "ABCD":
                correct = last.upper()
            found_opt = False
            for cell_text in cells[1:-1]:
                cell_text = cell_text.strip()
                if not cell_text:
                    continue
                m = opt_re.match(cell_text)
                if m:
                    opts[m.group(1).upper()] = m.group(2).strip()
                    found_opt = True
            if not found_opt:
                for cell_text in cells[1:-1]:
                    if cell_text.strip():
                        stem = cell_text.strip()
                        break
        missing = [L for L in "ABCD" if L not in opts]
        if not stem:
            errors.append(f"Question {num}: missing question text.")
        elif missing:
            errors.append(f"Question {num} ('{stem[:40]}...'): missing option(s) {', '.join(missing)}.")
        elif not correct:
            errors.append(f"Question {num} ('{stem[:40]}...'): no correct-answer letter found.")
        else:
            questions.append({
                "q": stem,
                "options": [opts["A"], opts["B"], opts["C"], opts["D"]],
                "correct_index": "ABCD".index(correct),
                "time": 30,
            })
    return questions, errors


def extract_questions_from_upload(uploaded_file):
    """Handles both the numbered-table Word format and the plain-text
    Q:/A)/Correct: format, plus .txt files. Returns (questions, errors)."""
    if uploaded_file.name.lower().endswith(".docx"):
        doc = Document(io.BytesIO(uploaded_file.read()))
        if doc.tables:
            # Try every table and keep whichever one yields the most valid questions
            best_q, best_err = [], ["No recognizable table or question text found in this file."]
            for table in doc.tables:
                qs, errs = parse_docx_table(table)
                if len(qs) > len(best_q):
                    best_q, best_err = qs, errs
            if best_q:
                return best_q, best_err
        # No usable table — fall back to reading it as plain Q:/A)/Correct: paragraphs
        paragraphs = [p.text for p in doc.paragraphs]
        return parse_questions_text(paragraphs)
    else:
        lines = uploaded_file.read().decode("utf-8", errors="ignore").split("\n")
        return parse_questions_text(lines)


# ============================================================
# STYLING
# ============================================================
st.set_page_config(page_title="Biopharmaceutics Live Quiz", page_icon="💊", layout="centered")
st.markdown("""
<style>
  html, body, [class*="css"] { font-family: 'IBM Plex Sans', sans-serif; }
  .quiz-card { background:#fff; border:1px solid #D6E2DC; border-radius:18px; padding:32px; }
  .eyebrow { font-family:monospace; font-size:12px; letter-spacing:.14em; text-transform:uppercase; color:#A66E1E; font-weight:600; }
  .big-code { font-family:monospace; font-size:44px; font-weight:700; color:#1F7A6C; text-align:center; margin:10px 0; }
  .qtext { font-size:22px; font-weight:600; text-align:center; margin:18px 0; }
  .board-row { display:flex; justify-content:space-between; padding:10px 14px; border-bottom:1px solid #D6E2DC; }
  .rank { font-family:monospace; font-weight:700; color:#A66E1E; }
  .score { font-family:monospace; font-weight:700; color:#1F7A6C; }

  @keyframes slideFadeIn {
    from { opacity: 0; transform: translateY(18px); }
    to   { opacity: 1; transform: translateY(0); }
  }
  @keyframes popIn {
    0%   { opacity: 0; transform: scale(0.85); }
    70%  { opacity: 1; transform: scale(1.03); }
    100% { opacity: 1; transform: scale(1); }
  }
  @keyframes pulseWarn {
    0%, 100% { color: #C0392B; }
    50%      { color: #7a231b; }
  }
  .qtext-anim {
    font-size:22px; font-weight:600; text-align:center; margin:18px 0;
    animation: slideFadeIn 0.55s ease-out;
  }
  .reveal-anim { animation: popIn 0.5s ease-out; }
  .board-row-anim { animation: slideFadeIn 0.4s ease-out; }
  .timer-normal { font-family:monospace; }
  .timer-urgent { font-family:monospace; animation: pulseWarn 0.8s infinite; }
</style>
""", unsafe_allow_html=True)

state = get_state()
role = st.query_params.get("role", "host")


# ============================================================
# HOST VIEW
# ============================================================
def render_host():
    if not require_host_password(state):
        return
    st_autorefresh(interval=1000, key="host_autorefresh")
    st.markdown('<div class="eyebrow">HOST · BIOPHARMACEUTICS LIVE QUIZ</div>', unsafe_allow_html=True)

    if state["status"] == "lobby":
        st.title("Waiting room")
        st.write("Share the **student link** with your class, then start once everyone's in.")
        st.code(f"{_base_url()}/?role=student", language=None)
        roster_items = list(state["roster"].values())
        st.write(f"**{len(roster_items)} student(s) joined**")
        if roster_items:
            st.write(", ".join(f"{r['name']} ({r['sid']})" for r in roster_items))
        else:
            st.caption("No one yet — waiting...")

        with st.expander(f"📋 Manage questions ({len(state['questions'])} currently)"):
            if state.get("using_sheet"):
                st.caption("✅ Questions are stored in your Google Sheet — changes here save there too, so they survive app restarts.")
                if st.button("🔄 Reload from Google Sheet"):
                    fresh, ferr = load_questions_from_sheet()
                    if fresh:
                        state["questions"] = fresh
                        st.success(f"Reloaded {len(fresh)} question(s) from the sheet.")
                    else:
                        st.warning(f"Couldn't reload: {ferr}")
                    st.rerun()
            else:
                st.warning(f"⚠️ Not connected to Google Sheets ({state.get('sheet_error', 'unknown reason')}). Using built-in defaults for now — changes here will be lost if the app restarts. See the setup guide.")
            for i, q in enumerate(state["questions"]):
                cols = st.columns([7, 1])
                correct_opt = q["options"][q["correct_index"]]
                cols[0].markdown(f"**Q{i+1}.** {q['q']}  \n*Correct: {correct_opt} · {q['time']}s*")
                if cols[1].button("Remove", key=f"delq_{i}"):
                    state["questions"].pop(i)
                    ok, err = save_questions_to_sheet(state["questions"])
                    if not ok and state.get("using_sheet"):
                        st.warning(f"Removed here, but couldn't save to the sheet: {err}")
                    st.rerun()

            st.write("---")
            st.write("**Upload questions from a Word or text file**")
            st.caption(
                "Two formats are supported — use whichever you already have:\n\n"
                "**1) Numbered table** (question #, stem, options A-D, correct letter column) — "
                "just export your existing quiz table as .docx.\n\n"
                "**2) Plain text**, one block per question, blank line between:\n"
                "Q: Your question text?\n"
                "A) First option\n"
                "B) Second option\n"
                "C) Third option\n"
                "D) Fourth option\n"
                "Correct: B\n"
                "Time: 30   (optional — defaults to 30 seconds if left out)"
            )
            uploaded = st.file_uploader("Upload .docx or .txt", type=["docx", "txt"], key="q_upload")
            if uploaded is not None:
                parsed, parse_errors = extract_questions_from_upload(uploaded)
                if parse_errors:
                    st.warning("Some questions couldn't be read:\n\n" + "\n".join(f"- {e}" for e in parse_errors))
                if parsed:
                    st.success(f"Found {len(parsed)} valid question(s) in this file.")
                    with st.container(border=True):
                        for i, q in enumerate(parsed):
                            st.write(f"**{i+1}.** {q['q']}")
                            st.caption(f"Correct: {q['options'][q['correct_index']]} · {q['time']}s")
                    import_mode = st.radio(
                        "How should these be added?",
                        ["Add to existing questions", "Replace all existing questions"],
                        key="import_mode",
                    )
                    if st.button("Import these questions", type="primary"):
                        if import_mode == "Replace all existing questions":
                            state["questions"] = parsed
                        else:
                            state["questions"].extend(parsed)
                        ok, err = save_questions_to_sheet(state["questions"])
                        if not ok and state.get("using_sheet"):
                            st.warning(f"Imported here, but couldn't save to the sheet: {err}")
                        st.rerun()
                else:
                    st.error("No valid questions found in this file. Check the format matches the example above.")

            st.write("---")
            st.write("**Or add one question manually**")
            new_q_text = st.text_input("Question text", key="new_q_text")
            opt_cols = st.columns(2)
            new_opts = []
            for j in range(4):
                new_opts.append(opt_cols[j % 2].text_input(f"Option {chr(65+j)}", key=f"new_opt_{j}"))
            new_correct = st.selectbox("Correct answer", options=[0, 1, 2, 3], format_func=lambda x: f"Option {chr(65+x)}", key="new_correct")
            new_time = st.number_input("Time limit (seconds)", min_value=5, max_value=120, value=30, step=5, key="new_time")
            if st.button("Add question"):
                if new_q_text.strip() and all(o.strip() for o in new_opts):
                    state["questions"].append({
                        "q": new_q_text.strip(),
                        "options": [o.strip() for o in new_opts],
                        "correct_index": new_correct,
                        "time": int(new_time),
                    })
                    ok, err = save_questions_to_sheet(state["questions"])
                    if not ok and state.get("using_sheet"):
                        st.warning(f"Added here, but couldn't save to the sheet: {err}")
                    for k in ["new_q_text", "new_opt_0", "new_opt_1", "new_opt_2", "new_opt_3"]:
                        if k in st.session_state:
                            del st.session_state[k]
                    st.rerun()
                else:
                    st.warning("Fill in the question text and all 4 options first.")

        if st.button("Start quiz", disabled=len(roster_items) == 0 or len(state["questions"]) == 0, type="primary"):
            state["status"] = "active"
            state["current_q"] = 0
            state["question_started_at"] = time.time()
            st.rerun()
        if st.button("Reset session"):
            reset_quiz(state)
            st.rerun()

    elif state["status"] == "active":
        q_idx = state["current_q"]
        q = state["questions"][q_idx]
        elapsed = time.time() - state["question_started_at"]
        remaining = max(0, q["time"] - elapsed)

        if "host_last_q" not in st.session_state:
            st.session_state.host_last_q = -1

        st.subheader(f"Question {q_idx+1} of {len(state['questions'])}")
        st.progress(min(1.0, remaining / q["time"]))
        timer_cls = "timer-urgent" if remaining <= 5 else "timer-normal"
        st.markdown(f'<div class="{timer_cls}"><h3>⏱ {int(remaining)+1}s</h3></div>', unsafe_allow_html=True)

        qtext_slot = st.empty()
        if st.session_state.host_last_q != q_idx:
            with qtext_slot.container():
                st.markdown(f'<div class="qtext-anim">{q["q"]}</div>', unsafe_allow_html=True)
                for i, opt in enumerate(q["options"]):
                    st.write(f"{chr(65+i)}. {opt}")
            st.session_state.host_last_q = q_idx

        qanswers = state["answers"].get(q_idx, {})
        st.caption(f"{len(qanswers)} of {len(state['roster'])} students have answered")
        with st.expander("Live responses", expanded=True):
            # answered students first (fastest at top), then everyone still waiting
            answered_rows = sorted(
                [(pid, a["elapsed"]) for pid, a in qanswers.items()],
                key=lambda x: x[1],
            )
            for pid, e in answered_rows:
                info = state["roster"].get(pid, {"name": "?", "sid": "?"})
                st.write(f"✅ {info['name']} ({info['sid']}) — answered in {e:.1f}s")
            waiting_pids = [pid for pid in state["roster"] if pid not in qanswers]
            for pid in waiting_pids:
                info = state["roster"].get(pid, {"name": "?", "sid": "?"})
                st.write(f"⏳ {info['name']} ({info['sid']}) — still answering...")

        if remaining <= 0:
            # anyone who never answered counts as a miss, then rank + score the round
            for pid in state["roster"]:
                if pid not in qanswers:
                    qanswers[pid] = {"choice": -1, "correct": False, "elapsed": q["time"]}
            state["answers"][q_idx] = qanswers
            rank_and_score_round(state, q_idx, q["time"])
            state["status"] = "reveal"
            state["reveal_started_at"] = time.time()
            st.rerun()

    elif state["status"] == "reveal":
        q_idx = state["current_q"]
        q = state["questions"][q_idx]

        if "host_last_reveal_q" not in st.session_state:
            st.session_state.host_last_reveal_q = -1

        reveal_slot = st.empty()
        if st.session_state.host_last_reveal_q != q_idx:
            with reveal_slot.container():
                st.subheader(f"Question {q_idx+1} — Answer")
                st.markdown(f'<div class="qtext-anim">{q["q"]}</div>', unsafe_allow_html=True)
                for i, opt in enumerate(q["options"]):
                    marker = "✅ " if i == q["correct_index"] else "▫️ "
                    st.write(f"{marker}{chr(65+i)}. {opt}")
                st.write("---")
                st.write("**Leaderboard so far**")
                for i, row in enumerate(leaderboard(state)[:8]):
                    st.markdown(f'<div class="board-row board-row-anim"><span class="rank">#{i+1}</span><span>{row["name"]} ({row["sid"]})</span><span class="score">{row["score"]}</span></div>', unsafe_allow_html=True)
            st.session_state.host_last_reveal_q = q_idx

        if time.time() - state["reveal_started_at"] >= REVEAL_SECONDS:
            nxt = state["current_q"] + 1
            if nxt >= len(state["questions"]):
                state["status"] = "finished"
            else:
                state["current_q"] = nxt
                state["status"] = "active"
                state["question_started_at"] = time.time()
            st.rerun()

    elif state["status"] == "finished":
        st.title("🏁 Quiz complete")
        st.write(f"{len(state['questions'])} questions · {len(state['roster'])} students")
        board = leaderboard(state)
        for i, row in enumerate(board):
            st.markdown(f'<div class="board-row"><span class="rank">#{i+1}</span><span>{row["name"]} ({row["sid"]})</span><span class="score">{row["score"]}</span></div>', unsafe_allow_html=True)

        timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H%M")

        st.write("---")
        st.subheader("📥 Save results")
        st.caption("This records exactly what each student selected and when, as it happened live — useful if a student disputes their score.")

        summary_rows = [{"Rank": i + 1, "Student Name": r["name"], "Student ID": r["sid"], "Total Score": r["score"]} for i, r in enumerate(board)]
        detail_rows = detailed_results(state)

        col1, col2 = st.columns(2)
        col1.download_button(
            "Download leaderboard (CSV)",
            data=rows_to_csv_bytes(summary_rows),
            file_name=f"quiz_leaderboard_{timestamp}.csv",
            mime="text/csv",
        )
        col2.download_button(
            "Download full answer log (CSV)",
            data=rows_to_csv_bytes(detail_rows),
            file_name=f"quiz_answers_{timestamp}.csv",
            mime="text/csv",
        )

        with st.expander("🔍 Review individual student answers"):
            student_options = {f"{info['name']} ({info['sid']})": pid for pid, info in state["roster"].items()}
            if student_options:
                pick = st.selectbox("Choose a student", list(student_options.keys()))
                chosen_pid = student_options[pick]
                for qi, q in enumerate(state["questions"]):
                    a = state["answers"].get(qi, {}).get(chosen_pid)
                    correct_letter = chr(65 + q["correct_index"])
                    if a is None:
                        sel = "No answer recorded"
                    elif a["choice"] == -1:
                        sel = "No answer (timed out)"
                    else:
                        sel = f"{chr(65+a['choice'])}. {q['options'][a['choice']]}"
                    ok = a.get("correct") if a else False
                    pts = a.get("score", 0) if a else 0
                    took = f"{a.get('elapsed', 0):.1f}s" if a else "—"
                    marker = "✅" if ok else "❌"
                    st.write(f"{marker} **Q{qi+1}.** {q['q']}")
                    st.caption(f"Selected: {sel} · Correct answer: {correct_letter}. {q['options'][q['correct_index']]} · Time: {took} · Points: {pts}")
            else:
                st.caption("No students joined this session.")

        st.write("---")
        if st.button("Start a new session"):
            reset_quiz(state)
            st.rerun()


# ============================================================
# STUDENT VIEW
# ============================================================
def render_student():
    if "pid" not in st.session_state:
        st.session_state.pid = str(uuid.uuid4())
    if "joined" not in st.session_state:
        st.session_state.joined = False

    pid = st.session_state.pid

    if not st.session_state.joined:
        st.markdown('<div class="eyebrow">BIOPHARMACEUTICS LIVE QUIZ</div>', unsafe_allow_html=True)
        st.title("Join the quiz")
        name = st.text_input("Your name", max_chars=24)
        sid = st.text_input("Your student ID", max_chars=24)
        if st.button("Join quiz", type="primary"):
            if name.strip() and sid.strip():
                state["roster"][pid] = {"name": name.strip(), "sid": sid.strip()}
                st.session_state.joined = True
                st.rerun()
            else:
                st.warning("Enter both your name and student ID — names can repeat, so the ID keeps your score yours.")
        return

    st_autorefresh(interval=700, key="student_autorefresh")
    my_info = state["roster"].get(pid, {"name": "you", "sid": ""})
    name = my_info["name"]

    if state["status"] == "lobby":
        st.title(f"Hi {name} 👋")
        st.write("Waiting for the host to start the quiz...")
        return

    if state["status"] == "finished":
        board = leaderboard(state)
        my_rank = next((i + 1 for i, r in enumerate(board) if r["pid"] == pid), None)
        my_score = next((r["score"] for r in board if r["pid"] == pid), 0)
        st.title(f"You finished #{my_rank or '-'} with {my_score} points")
        for i, row in enumerate(board):
            st.markdown(f'<div class="board-row"><span class="rank">#{i+1}</span><span>{row["name"]} ({row["sid"]})</span><span class="score">{row["score"]}</span></div>', unsafe_allow_html=True)
        return

    q_idx = state["current_q"]
    q = state["questions"][q_idx]

    if state["status"] == "active":
        elapsed = time.time() - state["question_started_at"]
        remaining = max(0, q["time"] - elapsed)
        already_answered = state["answers"].get(q_idx, {}).get(pid) is not None

        if "student_last_q" not in st.session_state:
            st.session_state.student_last_q = -1

        correct_so_far, attempted_so_far = my_progress(state, pid, q_idx - 1)
        if attempted_so_far > 0:
            st.caption(f"✅ {correct_so_far} of {attempted_so_far} correct so far")

        st.subheader(f"Question {q_idx+1} of {len(state['questions'])}")
        st.progress(min(1.0, remaining / q["time"]))
        timer_cls = "timer-urgent" if remaining <= 5 else "timer-normal"
        st.markdown(f'<div class="{timer_cls}"><h3>⏱ {int(remaining)+1}s</h3></div>', unsafe_allow_html=True)

        qtext_slot = st.empty()
        if st.session_state.student_last_q != q_idx:
            with qtext_slot.container():
                st.markdown(f'<div class="qtext-anim">{q["q"]}</div>', unsafe_allow_html=True)
            st.session_state.student_last_q = q_idx

        if already_answered:
            st.info("Answer locked — waiting for the round to end.")
        elif remaining <= 0:
            st.info("Time's up — waiting for the round to end.")
        else:
            cols = st.columns(2)
            for i, opt in enumerate(q["options"]):
                if cols[i % 2].button(f"{chr(65+i)}. {opt}", key=f"opt_{q_idx}_{i}", use_container_width=True):
                    e = min(q["time"], time.time() - state["question_started_at"])
                    correct = (i == q["correct_index"])
                    # Score is intentionally NOT set here — the host ranks every
                    # correct answer by speed once the timer ends, so ties and
                    # relative speed are only knowable after everyone's in.
                    state["answers"].setdefault(q_idx, {})[pid] = {"choice": i, "correct": correct, "elapsed": e}
                    st.rerun()

    elif state["status"] == "reveal":
        if "student_last_reveal_q" not in st.session_state:
            st.session_state.student_last_reveal_q = -1

        reveal_slot = st.empty()
        if st.session_state.student_last_reveal_q != q_idx:
            my_answer = state["answers"].get(q_idx, {}).get(pid)
            correct = my_answer["correct"] if my_answer else False
            score = my_answer.get("score", 0) if my_answer else 0
            correct_so_far, attempted_so_far = my_progress(state, pid, q_idx)
            with reveal_slot.container():
                st.caption(f"✅ {correct_so_far} of {attempted_so_far} correct so far")
                st.markdown(
                    f'<div class="reveal-anim"><h1>{"Correct! ✅" if correct else "Not quite ❌"}</h1></div>',
                    unsafe_allow_html=True,
                )
                st.write(f"You scored **{score}** points this round")
                for i, opt in enumerate(q["options"]):
                    marker = "✅ " if i == q["correct_index"] else "▫️ "
                    st.write(f"{marker}{chr(65+i)}. {opt}")
            st.session_state.student_last_reveal_q = q_idx


def _base_url():
    # Best-effort guess for display purposes only; the actual link is whatever
    # this app's deployed Streamlit Cloud URL is.
    return "https://<your-app-name>.streamlit.app"


if role == "student":
    render_student()
else:
    render_host()
